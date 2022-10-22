import streamlit as st
import streamlit.components.v1 as components 

# NLP Pkgs
import docuscospacy.corpus_analysis as ds

import pandas as pd
import plotly.express as px
from collections import Counter

import base64
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT

st.title("Explore single texts")

if 'corpus' not in st.session_state:
	st.session_state.corpus = ''
	
if 'docids' not in st.session_state:
	st.session_state.docids = ''

if 'doc_pos' not in st.session_state:
	st.session_state.doc_pos = ''

if 'doc_ds' not in st.session_state:
	st.session_state.doc_ds = ''

if 'dc_pos' not in st.session_state:
	st.session_state.dc_pos = ''

if 'dc_ds' not in st.session_state:
	st.session_state.dc_ds = ''

if 'html_pos' not in st.session_state:
	st.session_state.html_pos = ''

if 'html_ds' not in st.session_state:
	st.session_state.html_ds = ''

if 'html_str' not in st.session_state:
	st.session_state.html_str = ''
	
if 'doc_key' not in st.session_state:
	st.session_state.doc_key = ''

if 'tags' not in st.session_state:
	st.session_state.tags = []

if 'count_4' not in st.session_state:
	st.session_state.count_4 = 0

def increment_counter():
	st.session_state.count_4 += 1

if st.session_state.count_4 % 2 == 0:
    idx = 0
else:
    idx = 1
    
hex_highlights = ['#5fb7ca', '#e35be5', '#ffc701', '#fe5b05', '#cb7d60']
html_highlights = [' { background-color:#5fb7ca; }', ' { background-color:#e35be5; }', ' { background-color:#ffc701; }', ' { background-color:#fe5b05; }', ' { background-color:#cb7d60; }']

def html_build(tok, key, count_by="tag"):
    df = ds.tag_ruler(tok=tok, key=key, count_by=count_by)
    df['ws'] = df['Token'].str.extract(r'(\s+)$')
    df['Token'] = df['Token'].str.replace(r'(\s+)$', '')
    df.Token[df['Tag'] != 'Untagged'] = df['Token'].str.replace(r'^(.*?)$', '\\1</span>')
    df = df.iloc[:,[1,0,4]]
    df.fillna('', inplace=True)
    df.Tag[df['Tag'] != 'Untagged'] = df['Tag'].str.replace(r'^(.*?)$', '<span class="\\1">')
    df.Tag[df['Tag'] == 'Untagged'] = df['Tag'].str.replace('Untagged', '')
    df['Text'] = df['Tag'] + df['Token'] + df['ws']
    doc = ''.join(df['Text'].tolist())
    return(doc)

def doc_counts(doc_span, n_tokens, count_by='pos'):
    if count_by=='pos':
        df = Counter(doc_span[doc_span.Tag != 'Y'].Tag)
        df = pd.DataFrame.from_dict(df, orient='index').reset_index()
        df = df.rename(columns={'index':'Tag', 0:'AF'})
        df['RF'] = df.AF/n_tokens*100
        df.sort_values(by=['AF', 'Tag'], ascending=[False, True], inplace=True)
        df.reset_index(drop=True, inplace=True)
    elif count_by=='ds':
        df = Counter(doc_span.Tag)
        df = pd.DataFrame.from_dict(df, orient='index').reset_index()
        df = df.rename(columns={'index':'Tag', 0:'AF'})
        df = df[df.Tag != 'Untagged']
        df['RF'] = df.AF/n_tokens*100
        df.sort_values(by=['AF', 'Tag'], ascending=[False, True], inplace=True)
        df.reset_index(drop=True, inplace=True)
    return(df)

def lexdensity_plot(df, tag_list):
	plot_colors = hex_highlights[:len(tag_list)]
	hts = [200, 280, 340, 420, 500]
	ht = hts[len(tag_list)-1]
	df['X'] = (df.index + 1)/(len(df.index))
	df = df[df['Tag'].isin(tag_list)]
	df['Y'] = 1
	df_b = df.copy()
	df_b['Y'] = 0
	df = pd.concat([df, df_b], axis=0)
	plot = px.line(df, x='X', y='Y', color='Tag', line_group='X', color_discrete_sequence=plot_colors, facet_row='Tag', category_orders = {'Tag':tag_list})
	plot.update_yaxes(title_text='', showticklabels=False, range = [0,1], showgrid=False, mirror=True, showline=True, linecolor='black')
	plot.update_xaxes(title_text='', range = [0,1], tick0=.25, dtick=.25, tickformat=".0%", mirror=True, showline=True, linecolor='black')
	plot.for_each_annotation(lambda a: a.update(text=a.text.split("=")[-1]))
	plot.update_layout(showlegend=False, paper_bgcolor='white', plot_bgcolor='white', height=ht)
	return(plot)

def update_tags(html_state):
	tags = st.session_state.tags
	tags = ['.' + x for x in tags]
	highlights = html_highlights[:len(tags)]
	style_str = [''.join(x) for x in zip(tags, highlights)]
	style_str = ''.join(style_str)
	style_sheet_str = '<style>' + style_str + '</style>'
	st.session_state.html_str = style_sheet_str + html_state

def add_alt_chunk(doc: Document, html: str):
    package = doc.part.package
    partname = package.next_partname('/word/altChunk%d.html')
    alt_part = Part(partname, 'text/html', html.encode(), package)
    r_id = doc.part.relate_to(alt_part, RT.A_F_CHUNK)
    alt_chunk = OxmlElement('w:altChunk')
    alt_chunk.set(qn('r:id'), r_id)
    doc.element.body.sectPr.addprevious(alt_chunk)

if bool(isinstance(st.session_state.dc_pos, pd.DataFrame)) == True:
	st.write("Use the menus to select the tags you would like to highlight.")
	
	tag_radio = st.radio("Select tags to display:", ("Parts-of-Speech", "DocuScope"), index=idx, on_change=increment_counter, horizontal=True)

	if tag_radio == 'Parts-of-Speech':
		tag_list = st.multiselect('Select tags to highlight', st.session_state.tags_pos, on_change = update_tags(st.session_state.html_pos), key='tags')
		#html_str = st.session_state.html_pos
		tag_loc = st.session_state.doc_pos
		df = st.session_state.dc_pos
	else:
		tag_list = st.multiselect('Select tags to highlight', st.session_state.tags_ds, on_change = update_tags(st.session_state.html_ds), key='tags')
		#html_str = st.session_state.html_ds
		tag_loc = st.session_state.doc_ds
		df = st.session_state.dc_ds
	
	col1, col2= st.columns([1,1])
	with col1:
		if st.button("Tag Density Plot"):
			if len(tag_list) > 5:
				st.write(':no_entry_sign: You can only plot a maximum of 5 tags.')
			elif len(tag_list) == 0:
				st.write('There are no tags to plot.')
			else:
				st.plotly_chart(lexdensity_plot(tag_loc, tag_list), use_container_width=False)

	with col2:
		if st.button("Select a new document"):
			st.session_state.doc_pos = ''
			st.session_state.doc_ds = ''
			st.session_state.dc_pos = ''
			st.session_state.dc_ds = ''
			st.session_state.html_pos = ''
			st.session_state.html_ds = ''
			st.session_state.doc_key = ''
			del st.session_state['tags']
			st.experimental_rerun()
	
	with st.expander("Plot explanation"):
		st.write("""
				The plot(s) shows lines segment where tags occur in what might be called 'normalized text time.'
				For example, if you had a text 100 tokens long and a tag occurred at the 10th, 25th, and 60th token,
				the plot would show lines at 10%, 25%, and 60% along the x-axis.
				""")


	if len(tag_list) > 5:
		st.write(':no_entry_sign: You can only hightlight a maximum of 5 tags.')

	st.markdown(f"""
				###  {st.session_state.doc_key}
				""")

	components.html(st.session_state.html_str, height=500, scrolling=True)
	st.dataframe(df)
	
	if st.button("Download"):
		with st.spinner('Creating download link...'):
			doc_html = st.session_state.html_str.split('</style>')
			style_sheet_str = doc_html[0] + '</style>'
			html_str = doc_html[1]
			doc_html = '<!DOCTYPE html><html><head>' + style_sheet_str + '</head><body>' + html_str + '</body></html>'
			downloaded_file = Document()
			downloaded_file.add_heading(st.session_state.doc_key)
			downloaded_file.add_heading('Table of tag frequencies:', 3)
			#add counts table
			df['RF'] = df.RF.round(2)
			t = downloaded_file.add_table(df.shape[0]+1, df.shape[1])
			# add the header rows.
			for j in range(df.shape[-1]):
				t.cell(0,j).text = df.columns[j]
			# add the rest of the data frame
			for i in range(df.shape[0]):
				for j in range(df.shape[-1]):
					t.cell(i+1,j).text = str(df.values[i,j])
			t.style = 'LightShading-Accent1'
			downloaded_file.add_heading('Highlighted tags:', 3)
			downloaded_file.add_heading(', '.join(tag_list), 4)
			#add html
			add_alt_chunk(downloaded_file, doc_html)
			towrite = BytesIO()
			downloaded_file.save(towrite)
			towrite.seek(0)  # reset pointer
			b64 = base64.b64encode(towrite.read()).decode()
			st.success('Link generated!')
			linko= f'<a href="data:vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="document_tags.docx">Download Word file</a>'
			st.markdown(linko, unsafe_allow_html=True)

else:
	st.write("Use the menus to select the tags you would like to highlight.")
	
	doc_key = st.selectbox("Select document to view:", (sorted(st.session_state.docids)))
	if st.button("Process Document"):
		if st.session_state.corpus == '':
			st.write("It doesn't look like you've loaded a corpus yet.")
		else:
			doc_pos = ds.tag_ruler(st.session_state.corpus, doc_key, count_by='pos')
			doc_ds = ds.tag_ruler(st.session_state.corpus, doc_key, count_by='ds')
			doc_tokens = len(doc_pos.index)
			doc_words = len(doc_pos[doc_pos.Tag != 'Y'])
			dc_pos = doc_counts(doc_pos, doc_words, count_by='pos')
			dc_ds = doc_counts(doc_ds, doc_tokens, count_by='ds')
			html_pos = html_build(st.session_state.corpus, doc_key, count_by='pos')
			html_ds = html_build(st.session_state.corpus, doc_key, count_by='ds')
			st.session_state.doc_pos = doc_pos
			st.session_state.doc_ds = doc_ds
			st.session_state.dc_pos = dc_pos
			st.session_state.dc_ds = dc_ds
			st.session_state.html_pos = html_pos
			st.session_state.html_ds = html_ds
			st.session_state.doc_key = doc_key
			st.experimental_rerun()

# Copyright (C) 2025 David West Brown

# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at

#     http://www.apache.org/licenses/LICENSE-2.0

# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

# DataFrame formatting for display and export in Streamlit apps
# Plotting functions
# Data enrichement helpers

import base64
import docx
from docx.shared import RGBColor
from docx.shared import Pt
from io import BytesIO
import math
import numpy as np
import os
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import polars as pl
import streamlit as st
import zipfile


def get_streamlit_column_config(
        df: pl.DataFrame | pd.DataFrame
        ) -> dict:
    """
    Returns a column_config dictionary for st.dataframe based on column name patterns,
    including helpful tooltips for each column.
    Adjusts RF tooltips based on whether the table is token-based or tag-based.
    """
    # Detect if this is a tags-only table (no 'Token' or 'Token_*' column)
    tags_only = not any(col.startswith("Token") for col in df.columns)

    # Define tooltips for common columns
    tooltips = {
        "AF": "Absolute frequency (raw count)",
        "RF": (
            "Relative frequency (percent of tokens)"
            if tags_only else
            "Relative frequency (per million tokens)"
        ),
        "LL": "Log-likelihood (keyness statistic)",
        "LR": "Log ratio (effect size)",
        "Range": "Document range (proportion of docs containing item)",
        "PV": "p-value (statistical significance)",
        "MI": "Mutual information (association strength)",
        "AF_Ref": "Absolute frequency in reference corpus",
        "RF_Ref": (
            "Relative frequency in reference corpus (percent of tokens)"
            if tags_only else
            "Relative frequency in reference corpus (per million tokens)"
        ),
        "Range_Ref": "Document range in reference corpus",
    }

    config = {}
    for col in df.columns:
        # Find base name for tooltip matching (handles e.g. "RF_Ref")
        base = col
        if col.endswith("_Ref"):
            base = col
        elif "_" in col:
            base = col.split("_")[0]
        # Set format and help
        if col.startswith("AF"):
            config[col] = st.column_config.NumberColumn(
                format="%.0f",
                help=tooltips.get(col, tooltips.get(base, "Absolute frequency"))
            )
        elif col.startswith("RF"):
            config[col] = st.column_config.NumberColumn(
                format="%.2f",
                help=tooltips.get(col, tooltips.get(base, "Relative frequency"))
            )
        elif col.startswith("LL"):
            config[col] = st.column_config.NumberColumn(
                format="%.2f",
                help=tooltips.get(col, tooltips.get(base, "Log-likelihood"))
            )
        elif col.startswith("LR"):
            config[col] = st.column_config.NumberColumn(
                format="%.2f",
                help=tooltips.get(col, tooltips.get(base, "Log ratio"))
            )
        elif col.startswith("Range"):
            config[col] = st.column_config.NumberColumn(
                format="%.2f %%",
                help=tooltips.get(col, tooltips.get(base, "Document range"))
            )
        elif col.startswith("PV"):
            config[col] = st.column_config.NumberColumn(
                format="%.3f",
                help=tooltips.get(col, tooltips.get(base, "p-value"))
            )
        elif col.startswith("MI"):
            config[col] = st.column_config.NumberColumn(
                format="%.3f",
                help=tooltips.get(col, tooltips.get(base, "Mutual information"))
            )
    return config


def add_category_description(
        cat_counts: dict,
        session: dict = None,
        corpus_type: str = "target"  # "target" or "reference"
        ) -> pd.DataFrame:
    """
    Adds a 'Category Description' column to cat_df
    if the corpus is internal and a mapping exists.
    Also displays a documentation link button for
    internal corpora.
    """
    cat_df = pd.DataFrame(cat_counts.items(), columns=["Category", "Count"]).sort_values("Category")  # noqa: E501
    # Determine which session key to use
    db_key = f"{corpus_type}_db"
    target_db = session.get(db_key, [''])[0]
    if not target_db:
        return cat_df

    corpus_name = os.path.basename(target_db)

    # Documentation links for each corpus family
    doc_links = {
        "MICUSP": "https://browndw.github.io/docuscope-docs/datasets/micusp.html",
        "BAWE": "https://browndw.github.io/docuscope-docs/datasets/bawe.html",
        "ELSEVIER": "https://browndw.github.io/docuscope-docs/datasets/elsevier.html",
        "HAPE": "https://browndw.github.io/docuscope-docs/datasets/hape.html",
    }

    # Map corpus name to doc link by checking which family it belongs to
    doc_link = None
    for key in doc_links:
        if key in corpus_name:
            doc_link = doc_links[key]
            break

    mappings = {
        "A_MICUSP_mini": {
            "BIO": "Biology", "CEE": "Civil and Environmental Engineering", "CLS": "Classical Studies",  # noqa: E501
            "ECO": "Economics", "EDU": "Education", "ENG": "English", "HIS": "History",
            "IOE": "Industrial and Operational Engineering", "LIN": "Linguistics", "MEC": "Mechanical Engineering",  # noqa: E501
            "NRE": "Natural Resources", "NUR": "Nursing", "PHI": "Philosophy", "PHY": "Physics",  # noqa: E501
            "POL": "Political Science", "PSY": "Psychology", "SOC": "Sociology"
        },
        "B_MICUSP": {
            "BIO": "Biology", "CEE": "Civil and Environmental Engineering", "CLS": "Classical Studies",  # noqa: E501
            "ECO": "Economics", "EDU": "Education", "ENG": "English", "HIS": "History",
            "IOE": "Industrial and Operational Engineering", "LIN": "Linguistics", "MEC": "Mechanical Engineering",  # noqa: E501
            "NRE": "Natural Resources", "NUR": "Nursing", "PHI": "Philosophy", "PHY": "Physics",  # noqa: E501
            "POL": "Political Science", "PSY": "Psychology", "SOC": "Sociology"
        },
        "C_BAWE_mini": {
            "AH": "Arts and Humanities", "LS": "Life Sciences", "PS": "Physical Sciences", "SS": "Social Sciences"  # noqa: E501
        },
        "D_BAWE": {
            "AH": "Arts and Humanities", "LS": "Life Sciences", "PS": "Physical Sciences", "SS": "Social Sciences"  # noqa: E501
        },
        "E_ELSEVIER": {
            "ARTS": "Arts and Humanities", "BIOC": "Biochemistry, Genetics and Molecular Biology",  # noqa: E501
            "BUSI": "Business, Management and Accounting", "CENG": "Chemical Engineering", "CHEM": "Chemistry",  # noqa: E501
            "COMP": "Computer Science", "DECI": "Decision Sciences", "ECON": "Economics, Econometrics and Finance",  # noqa: E501
            "ENGI": "Engineering", "ENVI": "Environmental Science", "HEAL": "Health Professions",  # noqa: E501
            "IMMU": "Immunology and Microbiology", "MATE": "Material Science", "MATH": "Mathematics",  # noqa: E501
            "MEDI": "Medicine", "NEUR": "Neuroscience", "NURS": "Nursing", "PHYS": "Physics and Astronomy",  # noqa: E501
            "PSYC": "Psychology", "SOCI": "Social Sciences"
        },
        "G_MICUSP_by_level": {
            "G0": "Final Year Undergraduate", "G1": "First Year Graduate",
            "G2": "Second Year Graduate", "G3": "Third Year Graduate"
        },
        "F_MICUSP_by_paper": {},
        "H_HAPE_mini": {},
    }

    mapping = mappings.get(corpus_name)
    if mapping and "Category" in cat_df.columns:
        cat_df["Category Description"] = cat_df["Category"].map(mapping).fillna(cat_df["Category"])  # noqa: E501
        # Move 'Category Description' to the second column
        cols = list(cat_df.columns)
        if "Category Description" in cols:
            cols.insert(1, cols.pop(cols.index("Category Description")))
            cat_df = cat_df[cols]

    # Show documentation link if available
    if doc_link:
        parts = corpus_name.split('_')
        if len(parts) > 1:
            corpus_label = parts[1]
        else:
            corpus_label = corpus_name
        st.link_button(
            label=f"About {corpus_label} (documentation)",
            url=doc_link,
            icon=":material/info:"
        )

    return cat_df


def plot_download_link(
        fig: go.Figure,
        filename="plot.png",
        scale=2,
        button_text="Download high-res PNG"
        ) -> None:
    """
    Display a download link for a high-resolution PNG of a Plotly figure.

    Parameters
    ----------
    fig : plotly.graph_objs.Figure
        The Plotly figure to export.
    filename : str
        The filename for the downloaded PNG.
    scale : int or float
        The scale factor for the image resolution (default 2).
    button_text : str
        The text to display for the download link/button.

    Returns
    -------
    None
        Renders a download link in the Streamlit app.
    """
    # Export the figure to a PNG byte stream
    fig.update_xaxes(automargin=True)
    fig.update_yaxes(automargin=True)
    img_bytes = fig.to_image(format="png", scale=scale)
    b64 = base64.b64encode(img_bytes).decode()
    href = f'<a href="data:image/png;base64,{b64}" download="{filename}">{button_text}</a>'
    st.markdown(href, unsafe_allow_html=True)


def plot_tag_frequencies_bar(
        df: pl.DataFrame | pd.DataFrame
        ) -> go.Figure:
    """
    Plot a horizontal bar chart of tag frequencies.
    Expects columns: 'Tag' and 'RF' (relative frequency, as percent).
    """
    # Sort tags by frequency descending
    df_sorted = df.sort('RF', descending=True) if hasattr(df, 'sort') else df.sort_values('RF', ascending=True)  # noqa: E501
    # If using polars, convert to pandas for Plotly
    if hasattr(df_sorted, 'to_pandas'):
        df_sorted = df_sorted.to_pandas()

    min_height = 200  # Minimum plot height in pixels
    height = max(24 * len(df_sorted) + 40, min_height)

    fig = px.bar(
        df_sorted,
        x='RF',
        y='Tag',
        orientation='h',
        color_discrete_sequence=['#133955'],
        hover_data={'Tag': True, 'RF': ':.2f'},
        height=height,
    )

    fig.update_layout(
        showlegend=False,
        margin=dict(l=0, r=0, t=30, b=40),
        xaxis_title='Frequency (% of tokens)',
        yaxis_title=None,
        yaxis=dict(autorange="reversed", tickfont=dict(size=12)),
    )
    fig.update_traces(
        marker_line_width=0,
        hovertemplate="<b>Tag:</b> %{y}<br><b>RF:</b> %{x:.2f}%<extra></extra>"
    )
    return fig


def plot_compare_corpus_bar(
        df: pl.DataFrame | pd.DataFrame
        ) -> go.Figure:
    """
    Plot a horizontal bar chart comparing tag frequencies in two corpus parts.
    Expects columns: 'Tag', 'RF', 'RF_Ref'.
    Parameters
    ----------
    df : pl.DataFrame or pd.DataFrame
        DataFrame containing tag frequencies with columns 'Tag', 'RF', and 'RF_Ref'.
        'RF' is the target corpus frequency, 'RF_Ref' is the reference corpus frequency.
    Returns
    -------
    fig : plotly.graph_objects.Figure
    """
    # Prepare DataFrame
    df_plot = df.to_pandas() if hasattr(df, "to_pandas") else df.copy()
    df_plot = df_plot[["Tag", "RF", "RF_Ref"]].copy()
    df_plot["Mean"] = df_plot[["RF", "RF_Ref"]].mean(axis=1)
    df_plot.rename(
        columns={"RF": "Target", "RF_Ref": "Reference"},
        inplace=True
    )
    df_plot = pd.melt(
        df_plot,
        id_vars=['Tag', 'Mean'],
        value_vars=['Target', 'Reference'],
        var_name='Corpus',
        value_name='RF'
    )
    # Do not sort after this point!

    # Set tag order by descending mean
    tag_order = df_plot.groupby("Tag")["Mean"].mean().sort_values(ascending=False).index.tolist()  # noqa: E501
    corpus_order = ['Reference', 'Target']  # Target will be on top

    height = max(24 * len(tag_order) + 100, 400)

    fig = px.bar(
        df_plot,
        x="RF",
        y="Tag",
        color="Corpus",
        color_discrete_sequence=["#e67e22", "#133955"],
        orientation="h",
        category_orders={"Tag": tag_order, "Corpus": corpus_order},
        hover_data={"Tag": True, "RF": ':.2f', "Corpus": True},
        height=height,
        custom_data=["Corpus"],  # <-- This ensures correct mapping
    )

    fig.update_layout(
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="left",
            x=0
        ),
        legend_title_text='',
        margin=dict(l=0, r=0, t=30, b=0),
        xaxis_title='Frequency (% of tokens)',
        yaxis_title=None,
        bargap=0.1,
        bargroupgap=0.05,
        barmode='group',
    )
    fig.update_traces(
        marker_line_width=0,
        hovertemplate="<b>Tag:</b> %{y}<br><b>Corpus:</b> %{customdata[0]}<br><b>RF:</b> %{x:.2f}%<extra></extra>",  # noqa: E501
    )
    return fig


def plot_general_boxplot(
        df: pl.DataFrame | pd.DataFrame,
        tag_col='Tag',
        value_col='RF',
        color=None,
        palette=None
        ) -> go.Figure:
    """
    General boxplot for the corpus, colored by tag, with legend at bottom left,
    and boxes sorted by median (highest to lowest).
    Allows user to specify a custom HEX color, a dict, or a Plotly palette.
    """
    # Sort tags by median value (descending)
    medians = df.groupby(tag_col)[value_col].median().sort_values(ascending=False)
    tag_order = medians.index.tolist()

    # Color logic
    if isinstance(color, dict):
        color_map = color
    elif isinstance(color, str) and color.lower().startswith("#"):
        color_map = {cat: color for cat in tag_order}
    elif palette:
        palette_colors = palette if isinstance(palette, list) else px.colors.qualitative.Set1
        color_map = {cat: palette_colors[i % len(palette_colors)] for i, cat in enumerate(tag_order)}
    else:
        palette_colors = px.colors.qualitative.Set1
        color_map = {cat: palette_colors[i % len(palette_colors)] for i, cat in enumerate(tag_order)}

    # Compute summary stats for hover
    stats = (
        df.groupby(tag_col)[value_col]
        .agg(['mean', 'median', lambda s: s.quantile(0.75) - s.quantile(0.25), 'min', 'max'])
        .rename(columns={'mean': 'Mean', 'median': 'Median', '<lambda_0>': 'IQR', 'min': 'Min', 'max': 'Max'})
        .reset_index()
    )

    # Create boxplot
    fig = px.box(
        df,
        x=value_col,
        y=tag_col,
        color=tag_col,
        color_discrete_map=color_map,
        points=False,
        orientation='h',
        category_orders={tag_col: tag_order}
    )

    # Turn off default boxplot hover for all traces
    for trace in fig.data:
        if trace.type == "box":
            trace.hoverinfo = "skip"
            trace.hoveron = "boxes"

    # Overlay transparent bar for custom hover
    bar_df = stats.copy()
    bar_df['bar'] = bar_df['Max'] - bar_df['Min']
    bar_df['base'] = bar_df['Min']
    fig2 = px.bar(
        bar_df,
        y=tag_col,
        x='bar',
        base='base',
        orientation='h',
        color=tag_col,
        color_discrete_map=color_map,
        hover_data={
            'Mean': ':.2f',
            'Median': ':.2f',
            'IQR': ':.2f',
            'Min': ':.2f',
            'Max': ':.2f',
            'base': False,
            'bar': False,
        },
    ).update_traces(opacity=0.01,  # nearly invisible, but hoverable
                    hovertemplate="<b>%{y}</b><br>Min: %{customdata[3]:.2f}%<br>IQR: %{customdata[2]:.2f}%<br>Median: %{customdata[1]:.2f}%<br>Mean: %{customdata[0]:.2f}%<br>Max: %{customdata[4]:.2f}%<extra></extra>")

    # Add bar traces to boxplot
    for trace in fig2.data:
        fig.add_trace(trace)

    fig.update_layout(
        hovermode="closest",
        showlegend=False,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.25,
            xanchor="left",
            x=0
        ),
        legend_title_text='',
        margin=dict(l=0, r=0, t=30, b=0),
        height=100 * len(tag_order) + 120,
        xaxis_title='Frequency (per 100 tokens)',
        yaxis_title="Tag"
    )
    fig.update_yaxes(showticklabels=True, title=None, tickangle=0)
    fig.update_xaxes(title_text='Frequency (per 100 tokens)', range=[0, None])
    # Add vertical line at x=0, colored gray
    fig.add_vline(
        x=0,
        line_width=2,
        line_color="lightgray",
        layer="below"
    )
    return fig


def plot_grouped_boxplot(
        df,
        tag_col='Tag',
        value_col='RF',
        group_col='Group',
        color=None,
        palette=None
        ) -> go.Figure:
    """
    Boxplot comparing categories in subcorpora, grouped by tag and colored by group.
    Allows user to specify a custom HEX color or a Plotly palette.
    """
    tag_order = (
        df.groupby(tag_col)[value_col].median().sort_values(ascending=False).index.tolist()
    )
    group_order = sorted(df[group_col].unique())

    if isinstance(color, dict):
        color_discrete_map = color
        color_arg = group_col
    elif isinstance(color, str) and color.lower().startswith("#"):
        color_discrete_map = {cat: color for cat in df[group_col].unique()}
        color_arg = group_col
    else:
        color_discrete_map = None
        color_arg = group_col

    # Compute summary stats for hover (per tag+group)
    stats = (
        df.groupby([tag_col, group_col])[value_col]
        .agg(['mean', 'median', lambda s: s.quantile(0.75) - s.quantile(0.25), 'min', 'max'])
        .rename(columns={'mean': 'Mean', 'median': 'Median', '<lambda_0>': 'IQR', 'min': 'Min', 'max': 'Max'})
        .reset_index()
    )

    # Create boxplot
    fig = px.box(
        df,
        y=tag_col,
        x=value_col,
        color=color_arg,
        color_discrete_map=color_discrete_map,
        color_discrete_sequence=palette if palette else None,
        points=False,
        orientation='h',
        category_orders={
            tag_col: tag_order,
            group_col: group_order
        },
        boxmode="group"
    )

    # Turn off default boxplot hover for all traces
    for trace in fig.data:
        if trace.type == "box":
            trace.hoverinfo = "skip"
            trace.hoveron = "boxes"

    # Overlay transparent bar for custom hover (per tag+group)
    bar_df = stats.copy()
    bar_df['bar'] = bar_df['Max'] - bar_df['Min']
    bar_df['base'] = bar_df['Min']
    fig2 = px.bar(
        bar_df,
        y=tag_col,
        x='bar',
        base='base',
        color=group_col,
        orientation='h',
        color_discrete_map=color_discrete_map,
        category_orders={group_col: group_order, tag_col: tag_order},
        hover_data={
            'Mean': ':.2f',
            'Median': ':.2f',
            'IQR': ':.2f',
            'Min': ':.2f',
            'Max': ':.2f',
            'base': False,
            'bar': False,
            tag_col: False,
            group_col: False,
        },
    ).update_traces(
        opacity=0.01,  # nearly invisible, but hoverable
        hovertemplate="<b>%{y} | %{customdata[5]}</b><br>Min: %{customdata[3]:.2f}%<br>IQR: %{customdata[2]:.2f}%<br>Median: %{customdata[1]:.2f}%<br>Mean: %{customdata[0]:.2f}%<br>Max: %{customdata[4]:.2f}%<extra></extra>"
    )

    # Add bar traces to boxplot
    for trace in fig2.data:
        trace.showlegend = False  # Hide bar overlay from legend
        fig.add_trace(trace)

    fig.update_layout(
        hovermode="closest",
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="top",
            y=-0.25,  # Move legend well below the plot
            xanchor="left",
            x=0
        ),
        legend_title_text='',
        margin=dict(l=0, r=0, t=30, b=0),
        height=40 * len(tag_order) + 120,
        xaxis_title='Frequency (per 100 tokens)',
        yaxis_title=None
    )
    fig.update_yaxes(showticklabels=True, title=None)
    fig.update_xaxes(title_text='Frequency (per 100 tokens)', range=[0, None])

    # Add vertical line at x=0, colored gray
    fig.add_vline(
        x=0,
        line_width=2,
        line_color="gray",
        layer="below"
    )
    return fig


def plot_scatter(
        df: pl.DataFrame,
        x_col: str,
        y_col: str,
        color=None,
        trendline: bool = False
        ) -> go.Figure:
    """
    Simple scatterplot for two variables, with optional color support and trendline.
    """
    x_label = x_col + ' (per 100 tokens)'
    y_label = y_col + ' (per 100 tokens)'

    # Determine color
    if isinstance(color, dict):
        color_val = list(color.values())[0] if color else '#133955'
    elif isinstance(color, str) and color.lower().startswith("#"):
        color_val = color
    else:
        color_val = '#133955'

    # Axis scaling and ticks
    x_max = df[x_col].max() if not df.empty else 1
    y_max = df[y_col].max() if not df.empty else 1
    axis_max = max(x_max, y_max)
    axis_max = axis_max * 1.05 if axis_max > 0 else 1

    # --- Tick calculation: min 4, max 8, multiples of 2.5 ---
    def get_tick_interval(axis_max):
        candidates = [0.5, 1, 2.5, 5, 10, 25, 50, 100]
        for interval in candidates:
            n_ticks = int(axis_max // interval) + 1
            if 4 <= n_ticks <= 8:
                return interval
        for interval in reversed(candidates):
            n_ticks = int(axis_max // interval) + 1
            if n_ticks >= 4:
                return interval
        return None  # fallback: let Plotly decide

    tick_interval = get_tick_interval(axis_max)
    if tick_interval:
        axis_max = math.ceil(axis_max / tick_interval) * tick_interval
        n_ticks = int(axis_max // tick_interval) + 1
        tickvals = [round(i * tick_interval, 2) for i in range(n_ticks)]
        ticktext = [str(v) for v in tickvals]
        xaxis_kwargs = dict(
            range=[0, axis_max],
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black',
            tickvals=tickvals,
            ticktext=ticktext
        )
        yaxis_kwargs = dict(
            range=[0, axis_max],
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black',
            tickvals=tickvals,
            ticktext=ticktext
        )
    else:
        xaxis_kwargs = dict(
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black'
        )
        yaxis_kwargs = dict(
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black'
        )

    fig = go.Figure()

    # All points
    fig.add_trace(go.Scatter(
        x=df[x_col],
        y=df[y_col],
        mode='markers',
        marker=dict(
            color=color_val,
            size=8,
            opacity=0.75,
            line=dict(width=0)
        ),
        name="All Points",
        text=df['doc_id'] if 'doc_id' in df.columns else None,
        hovertemplate=(
            "<b>doc_id:</b> %{text}<br>"
            f"<b>{x_col}:</b> %{{x:.2f}}%<br>"
            f"<b>{y_col}:</b> %{{y:.2f}}%<extra></extra>"
        ) if 'doc_id' in df.columns else (
            f"<b>{x_col}:</b> %{{x:.2f}}%<br>"
            f"<b>{y_col}:</b> %{{y:.2f}}%<extra></extra>"
        ),
        showlegend=False
    ))

    # Optional: Add trendline for all points
    if trendline and not df.empty and len(df) > 1:
        x = df[x_col]
        y = df[y_col]
        coeffs = np.polyfit(x, y, 1)
        x_fit = np.array([0, axis_max])
        y_fit = coeffs[0] * x_fit + coeffs[1]
        fig.add_trace(go.Scatter(
            x=x_fit,
            y=y_fit,
            mode='lines',
            line=dict(color='tomato', width=2, dash='dash'),
            name="Linear fit",
            showlegend=False
        ))

    fig.update_xaxes(**xaxis_kwargs)
    fig.update_yaxes(**yaxis_kwargs)
    fig.update_layout(
        showlegend=False,
        margin=dict(l=0, r=0, t=30, b=40),
        xaxis_title=x_label or x_col,
        yaxis_title=y_label or y_col,
        height=500,
        width=500
    )
    return fig


def plot_scatter_highlight(
        df: pl.DataFrame,
        x_col: str,
        y_col: str,
        group_col: str,
        selected_groups: list = None,
        color=None,
        trendline=None
        ) -> go.Figure:
    """
    Scatterplot with optional group highlighting and user-defined colors.
    Highlighted points are plotted on top of non-highlighted points.
    """
    x_label = x_col + ' (per 100 tokens)'
    y_label = y_col + ' (per 100 tokens)'
    df = df.copy()
    if 'Highlight' in df.columns:
        df = df.drop(columns=['Highlight'])
    df['Highlight'] = True
    if selected_groups:
        df['Highlight'] = df[group_col].apply(lambda g: g in selected_groups)
    else:
        df['Highlight'] = True

    # Color logic
    if isinstance(color, dict):
        highlight_color = color.get('Highlight', '#133955')
        non_highlight_color = color.get('Non-Highlight', 'lightgray')
    elif isinstance(color, str) and color.lower().startswith("#"):
        highlight_color = color
        non_highlight_color = 'lightgray'
    else:
        highlight_color = '#133955'
        non_highlight_color = 'lightgray'

    # Axis scaling and ticks
    x_max = df[x_col].max() if not df.empty else 1
    y_max = df[y_col].max() if not df.empty else 1
    axis_max = max(x_max, y_max)
    axis_max = axis_max * 1.05 if axis_max > 0 else 1

    # --- Tick calculation: min 4, max 8, multiples of 2.5 ---
    def get_tick_interval(axis_max):
        candidates = [0.5, 1, 2.5, 5, 10, 25, 50, 100]
        for interval in candidates:
            n_ticks = int(axis_max // interval) + 1
            if 4 <= n_ticks <= 8:
                return interval
        for interval in reversed(candidates):
            n_ticks = int(axis_max // interval) + 1
            if n_ticks >= 4:
                return interval
        return None  # fallback: let Plotly decide

    tick_interval = get_tick_interval(axis_max)
    if tick_interval:
        axis_max = math.ceil(axis_max / tick_interval) * tick_interval
        n_ticks = int(axis_max // tick_interval) + 1
        tickvals = [round(i * tick_interval, 2) for i in range(n_ticks)]
        ticktext = [str(v) for v in tickvals]
        xaxis_kwargs = dict(
            range=[0, axis_max],
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black',
            tickvals=tickvals,
            ticktext=ticktext
        )
        yaxis_kwargs = dict(
            range=[0, axis_max],
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black',
            tickvals=tickvals,
            ticktext=ticktext
        )
    else:
        xaxis_kwargs = dict(
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black'
        )
        yaxis_kwargs = dict(
            showgrid=False,
            zeroline=True,
            zerolinewidth=2,
            zerolinecolor='black'
        )

    # Split data
    df_non_highlight = df[df['Highlight'] == False]  # noqa: E712
    df_highlight = df[df['Highlight'] == True]  # noqa: E712

    fig = go.Figure()

    # Non-highlighted points (bottom layer)
    if not df_non_highlight.empty:
        fig.add_trace(go.Scatter(
            x=df_non_highlight[x_col],
            y=df_non_highlight[y_col],
            mode='markers',
            marker=dict(
                color=non_highlight_color,
                size=8,
                opacity=0.5,
                line=dict(width=0)
            ),
            name="Non-Highlight",
            text=df_non_highlight[group_col] if group_col in df_non_highlight.columns else None,
            hovertemplate=(
                f"<b>{group_col}:</b> %{{text}}<br>"
                f"<b>{x_col}:</b> %{{x:.2f}}%<br>"
                f"<b>{y_col}:</b> %{{y:.2f}}%<extra></extra>"
            ) if group_col in df_non_highlight.columns else None,
            showlegend=False
        ))

    # Highlighted points (top layer)
    if not df_highlight.empty:
        fig.add_trace(go.Scatter(
            x=df_highlight[x_col],
            y=df_highlight[y_col],
            mode='markers',
            marker=dict(
                color=highlight_color,
                size=8,
                opacity=0.85,
                line=dict(width=1, color='black')
            ),
            name="Highlight",
            text=df_highlight[group_col] if group_col in df_highlight.columns else None,
            hovertemplate=(
                f"<b>{group_col}:</b> %{{text}}<br>"
                f"<b>{x_col}:</b> %{{x:.2f}}%<br>"
                f"<b>{y_col}:</b> %{{y:.2f}}%<extra></extra>"
            ) if group_col in df_highlight.columns else None,
            showlegend=False
        ))

    # Optional: Add trendline for highlighted points only
    if trendline and not df_highlight.empty and len(df_highlight) > 1:
        x = df_highlight[x_col]
        y = df_highlight[y_col]
        coeffs = np.polyfit(x, y, 1)
        x_fit = np.array([0, axis_max])
        y_fit = coeffs[0] * x_fit + coeffs[1]
        fig.add_trace(go.Scatter(
            x=x_fit,
            y=y_fit,
            mode='lines',
            line=dict(color='tomato', width=2, dash='dash'),
            name="Linear fit",
            showlegend=False
        ))

    fig.update_xaxes(**xaxis_kwargs)
    fig.update_yaxes(**yaxis_kwargs)
    fig.update_layout(
        showlegend=False,
        margin=dict(l=0, r=0, t=30, b=40),
        xaxis_title=x_label or x_col,
        yaxis_title=y_label or y_col,
        height=500,
        width=500
    )
    return fig


def plot_pca_scatter_highlight(
        df: pl.DataFrame,
        x_col: str,
        y_col: str,
        group_col: str,
        selected_groups: list = None,
        x_label: str = None,
        y_label: str = None
        ) -> go.Figure:
    """
    Create a scatter plot for PCA results with optional highlighting of groups.
    Highlighted points are plotted on top of non-highlighted points.
    """
    # Convert to pandas if needed
    if hasattr(df, "to_pandas"):
        df = df.to_pandas()
    else:
        df = df.copy()

    # Drop 'Highlight' if present, then copy to avoid SettingWithCopyWarning
    if 'Highlight' in df.columns:
        df = df.drop(columns=['Highlight']).copy()
    else:
        df = df.copy()

    df['Highlight'] = True
    if selected_groups:
        df['Highlight'] = df[group_col].apply(lambda g: g in selected_groups)
    else:
        df['Highlight'] = True

    # Color logic
    highlight_color = '#133955'
    non_highlight_color = 'lightgray'

    # Find max absolute value for axis normalization
    max_abs = max(
        abs(df[x_col].min()), abs(df[x_col].max()),
        abs(df[y_col].min()), abs(df[y_col].max())
    )
    max_abs = (int(max_abs) + 1) if max_abs == int(max_abs) else round(max_abs + 0.5, 2)

    # Split data
    df_non_highlight = df[df['Highlight'] == False]
    df_highlight = df[df['Highlight'] == True]

    fig = go.Figure()

    # Non-highlighted points (bottom layer)
    if not df_non_highlight.empty:
        fig.add_trace(go.Scatter(
            x=df_non_highlight[x_col],
            y=df_non_highlight[y_col],
            mode='markers',
            marker=dict(
                color=non_highlight_color,
                size=8,
                opacity=0.5,
                line=dict(width=0)
            ),
            name="Non-Highlight",
            text=df_non_highlight[group_col] if group_col in df_non_highlight.columns else None,
            hovertemplate=(
                f"<b>{group_col}:</b> %{{text}}<br>"
                f"<b>{x_col}:</b> %{{x:.2f}}<br>"
                f"<b>{y_col}:</b> %{{y:.2f}}<extra></extra>"
            ) if group_col in df_non_highlight.columns else None,
            showlegend=False
        ))

    # Highlighted points (top layer)
    if not df_highlight.empty:
        fig.add_trace(go.Scatter(
            x=df_highlight[x_col],
            y=df_highlight[y_col],
            mode='markers',
            marker=dict(
                color=highlight_color,
                size=8,
                opacity=0.85,
                line=dict(width=1, color='black')
            ),
            name="Highlight",
            text=df_highlight[group_col] if group_col in df_highlight.columns else None,
            hovertemplate=(
                f"<b>{group_col}:</b> %{{text}}<br>"
                f"<b>{x_col}:</b> %{{x:.2f}}<br>"
                f"<b>{y_col}:</b> %{{y:.2f}}<extra></extra>"
            ) if group_col in df_highlight.columns else None,
            showlegend=False
        ))

    # Add zero axes
    fig.add_shape(type="line",
                  x0=0, x1=0,
                  y0=-max_abs, y1=max_abs,
                  line=dict(color="black", width=1, dash="dash"),
                  layer="below")
    fig.add_shape(type="line",
                  x0=-max_abs, x1=max_abs,
                  y0=0, y1=0,
                  line=dict(color="black", width=1, dash="dash"),
                  layer="below")

    fig.update_layout(
        showlegend=False,
        margin=dict(l=0, r=0, t=30, b=40),
        xaxis_title=x_label or x_col,
        yaxis_title=y_label or y_col,
        height=500,
        width=500
    )
    fig.update_xaxes(showgrid=False, range=[-max_abs, max_abs], zeroline=False)
    fig.update_yaxes(showgrid=False, range=[-max_abs, max_abs], zeroline=False)
    return fig


def plot_pca_variable_contrib_bar(
        contrib_1_plot,
        contrib_2_plot,
        pc1_label="PC1",
        pc2_label="PC2",
        sort_by=None
        ) -> go.Figure:
    """
    Create a horizontal bar plot comparing variable contributions
    to two principal components (PC1 and PC2).
    Parameters
    ----------
    contrib_1_plot : pd.DataFrame
        DataFrame containing contributions for PC1.
        Must have columns: 'Tag', 'Contribution'.
    contrib_2_plot : pd.DataFrame
        DataFrame containing contributions for PC2.
        Must have columns: 'Tag', 'Contribution'.
    pc1_label : str
        Label for PC1, default is "PC1".
    pc2_label : str
        Label for PC2, default is "PC2".
    sort_by : str, optional
        If provided, sort the bars by this PC label.
        If None, sort by PC1.
    Returns
    -------
    fig : plotly.graph_objects.Figure
        The resulting bar plot figure.
    """
    # Merge on Tag for alignment
    merged = contrib_1_plot.merge(
        contrib_2_plot, on="Tag", how="outer", suffixes=(f"_{pc1_label}", f"_{pc2_label}")
    ).fillna(0)

    # Get column names for contributions
    col_pc1 = merged.columns[1]
    col_pc2 = merged.columns[2]

    # Calculate mean absolute contributions
    mean_pc1 = merged[col_pc1].abs().mean()
    mean_pc2 = merged[col_pc2].abs().mean()

    # Decide which PC to sort by
    if sort_by == pc2_label:
        sort_col = col_pc2
        main_col = col_pc2
        mean_main = mean_pc2
        other_col = col_pc1
        mean_other = mean_pc1
    else:
        sort_col = col_pc1
        main_col = col_pc1
        mean_main = mean_pc1
        other_col = col_pc2
        mean_other = mean_pc2

    merged = merged.sort_values(by=sort_col, ascending=True)

    # Assign color and opacity for each bar
    colors_main = []
    opacities_main = []
    colors_other = []
    opacities_other = []

    for _, row in merged.iterrows():
        # Main (sorted-by) PC
        if abs(row[main_col]) > mean_main:
            colors_main.append("#133955")  # dark blue
            opacities_main.append(1.0)
        else:
            colors_main.append("#216495")  # light blue
            opacities_main.append(0.6)
        # Other PC always gray
        colors_other.append("#FFFFFF")  # white
        opacities_other.append(0.4)

    # Plot bars: main PC first, then other PC
    fig = go.Figure()
    # Main PC bars
    fig.add_trace(go.Bar(
        y=merged["Tag"],
        x=merged[main_col],
        name=sort_by if sort_by else pc1_label,
        orientation='h',
        marker_color=colors_main,
        opacity=1.0,
        hovertemplate=(
            f"<b>{sort_by if sort_by else pc1_label}</b><br>"
            "Variable: %{y}<br>"
            "Contribution: %{x:.2%}<extra></extra>"
        ),
        marker=dict(opacity=opacities_main)
    ))
    # Other PC bars
    fig.add_trace(go.Bar(
        y=merged["Tag"],
        x=merged[other_col],
        name=pc2_label if main_col == col_pc1 else pc1_label,
        orientation='h',
        marker_color=colors_other,
        opacity=1.0,
        hovertemplate=(
            f"<b>{pc2_label if main_col == col_pc1 else pc1_label}</b><br>"
            "Variable: %{y}<br>"
            "Contribution: %{x:.2%}<extra></extra>"
        ),
        marker=dict(opacity=opacities_other)
    ))

    # Add vertical lines for mean absolute contributions (main and other PC)
    for mean_val in [mean_main, -mean_main, mean_other, -mean_other]:
        fig.add_vline(
            x=mean_val,
            line=dict(color="tomato", width=2, dash="dot"),
            annotation_text="|mean|",
            annotation_position="top",
            opacity=0.7
        )

    # Set tick labels every 5% (0.05), covering the full range
    min_val = min(merged[col_pc1].min(), merged[col_pc2].min())
    max_val = max(merged[col_pc1].max(), merged[col_pc2].max())
    tick_start = (int(min_val * 20) - 1) / 20  # round down to nearest 0.05
    tick_end = (int(max_val * 20) + 1) / 20    # round up to nearest 0.05
    tickvals = [x / 100 for x in range(int(tick_start * 100), int(tick_end * 100) + 1, 5)]
    ticktext = [f"{abs(x)*100:.0f}%" for x in tickvals]

    fig.update_layout(
        barmode='group',
        height=30 * len(merged) + 100,
        margin=dict(l=0, r=0, t=30, b=40),
        xaxis_title="Contribution",
        yaxis_title="Variable",
        xaxis=dict(
            tickvals=tickvals,
            ticktext=ticktext,
            showgrid=True,
            gridcolor='lightgray',
            gridwidth=1
        ),
        showlegend=False,
    )
    return fig


# Functions for handling data tables
def convert_to_excel(
        df: pl.DataFrame
        ) -> bytes:
    """
    Convert a DataFrame to an Excel file in memory.

    Parameters
    ----------
    df : pandas.DataFrame
        The DataFrame to be converted to Excel format.

    Returns
    -------
    bytes
        The Excel file as a bytes object.
    """
    output = BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    df.to_excel(writer, index=False, header=True)
    writer.close()
    processed_data = output.getvalue()
    return processed_data


def add_alt_chunk(doc: docx.Document,
                  html: str) -> None:
    """
    Add an HTML altChunk to a Word document.

    Parameters
    ----------
    doc : docx.Document
        The Word document to which the altChunk will be added.
    html : str
        The HTML string to embed as an altChunk.

    Returns
    -------
    None
    """
    package = doc.part.package
    partname = package.next_partname('/word/altChunk%d.html')
    alt_part = docx.opc.part.Part(
        partname,
        'text/html',
        html.encode(),
        package
        )
    r_id = doc.part.relate_to(
        alt_part,
        docx.opc.constants.RELATIONSHIP_TYPE.A_F_CHUNK
        )
    alt_chunk = docx.oxml.OxmlElement('w:altChunk')
    alt_chunk.set(docx.oxml.ns.qn('r:id'), r_id)
    doc.element.body.sectPr.addprevious(alt_chunk)


def convert_to_word(
        html_string: str,
        tag_html: str,
        doc_key: str,
        tag_counts: pd.DataFrame
        ) -> bytes:
    """
    Convert HTML content and tag counts into a Word document and
    return it as bytes.

    Parameters
    ----------
    html_string : str
        The HTML string containing the document content and styles.
    tag_html : str
        The HTML string representing highlighted tags to embed in the document.
    doc_key : str
        The document key or title to use in the Word file.
    tag_counts : pd.DataFrame
        A DataFrame containing tag frequency information to include as a table.

    Returns
    -------
    bytes
        The generated Word document as a bytes object.
    """
    doc_html = html_string.split('</style>')
    style_sheet_str = doc_html[0] + '</style>'
    html_str = doc_html[1]
    doc_html = ('<!DOCTYPE html><html><head>' + style_sheet_str +
                '</head><body>' + tag_html +
                '<br><br>' + html_str +
                '</body></html>')
    download_file = docx.Document()
    title = download_file.add_heading(doc_key)
    title.style.font.color.rgb = RGBColor(0, 0, 0)
    heading = download_file.add_heading('Table of tag frequencies:', 3)
    heading.style.font.color.rgb = RGBColor(0, 0, 0)
    # add counts table
    tag_counts['RF'] = tag_counts.RF.round(2)
    t = download_file.add_table(tag_counts.shape[0]+1, tag_counts.shape[1])
    # add the header rows.
    for j in range(tag_counts.shape[-1]):
        t.cell(0, j).text = tag_counts.columns[j]
    # add the rest of the data frame
    for i in range(tag_counts.shape[0]):
        for j in range(tag_counts.shape[-1]):
            t.cell(i + 1, j).text = str(tag_counts.values[i, j])
    t.style = 'LightList'

    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
                    font.name = "Arial"

    download_file.add_heading('Highlighted tags:', 3)
    # add html
    add_alt_chunk(download_file, doc_html)
    output = BytesIO()
    download_file.save(output)
    processed_data = output.getvalue()
    return processed_data


def convert_corpus_to_zip(
        session_id: str,
        corpus_type: str,
        file_type="parquet"
        ) -> bytes:
    """
    Convert all tables in a corpus to a ZIP archive of Parquet or CSV files.

    Parameters
    ----------
    session_id : str
        The session ID for which the corpus is to be exported.
    corpus_type : str
        The type of corpus to export (e.g., 'target' or 'reference').
    file_type : str, optional
        The file format for export: 'parquet' (default) or 'csv'.

    Returns
    -------
    bytes
        The ZIP archive as a bytes object, with each table as a file.
    """
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as file_zip:
        for table in st.session_state[session_id][corpus_type]:
            _df = st.session_state[session_id][corpus_type][table]
            if file_type == "parquet":
                _df = _df.to_pandas().to_parquet()
                file_zip.writestr(table + ".parquet", _df)
            else:
                _df = _df.to_pandas().to_csv()
                file_zip.writestr(table + ".csv", _df)
    processed_data = zip_buf.getvalue()
    return processed_data


def convert_to_zip(
        tok_pl: pl.DataFrame,
        tagset: str
        ) -> bytes:
    """
    Convert tokenized corpus data to a ZIP archive of tagged text files.

    Parameters
    ----------
    tok_pl : pl.DataFrame
        The Polars DataFrame containing tokenized corpus data.
    tagset : str
        The tagset to use for tagging ('pos' or 'ds').

    Returns
    -------
    bytes
        The ZIP archive as a bytes object, with each document
        as a tagged text file.
    """
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as file_zip:
        for id in tok_pl.get_column("doc_id").unique().to_list():
            if tagset == "pos":
                df = (
                    tok_pl
                    .filter(pl.col("doc_id") == id)
                    .group_by(
                        ["pos_id", "pos_tag"], maintain_order=True
                        )
                    .agg(pl.col("token").str.concat(""))
                    .with_columns(
                        pl.col("token").str.strip_chars()
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(" ", "_")
                        )
                    .with_columns(
                        pl.when(pl.col("pos_tag") == "Y")
                        .then(pl.col("pos_tag").str.replace(
                            "Y", "", literal=True
                            ))
                        .when(pl.col("pos_tag") == "FU")
                        .then(pl.col("pos_tag").str.replace(
                            "FU", "", literal=True
                            ))
                        .otherwise(pl.col("pos_tag")))
                    .with_columns(
                        pl.concat_str(
                            pl.col("token"), pl.lit("|"), pl.col("pos_tag")
                            )
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(r"\|$", "")
                        )
                    )
            else:
                df = (
                    tok_pl
                    .filter(pl.col("doc_id") == id)
                    .group_by(["ds_id", "ds_tag"], maintain_order=True)
                    .agg(pl.col("token").str.concat(""))
                    .with_columns(
                        pl.col("token").str.strip_chars()
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(" ", "_")
                        )
                    .with_columns(
                        pl.when(
                            pl.col("ds_tag") == "Untagged"
                            )
                        .then(
                            pl.col("ds_tag").str.replace(
                                "Untagged",
                                "",
                                literal=True)
                            )
                        .otherwise(pl.col("ds_tag")))
                    .with_columns(
                        pl.concat_str(
                            pl.col("token"),
                            pl.lit("|"),
                            pl.col("ds_tag")
                            )
                        )
                    .with_columns(
                        pl.col("token").str.replace_all(r"\|$", "")
                        )
                    )
            doc = " ".join(df.get_column("token").to_list())
            file_zip.writestr(id + "_tagged" + ".txt", doc)
    processed_data = zip_buf.getvalue()

    return processed_data
